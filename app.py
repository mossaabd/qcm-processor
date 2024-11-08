# app.py
from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
import os
from docx import Document
import re
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize

# Initialize Flask app
app = Flask(__name__)

# Configure upload folder
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'

# Create folders if they don't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# Download required NLTK data
try:
    nltk.download('punkt', quiet=True)
    nltk.download('averaged_perceptron_tagger', quiet=True)
    nltk.download('punkt_tab', quiet=True)
except:
    print("Error downloading NLTK data")

class QuestionAnalyzer:
    def __init__(self):
        try:
            nltk.download('punkt', quiet=True)
            nltk.download('averaged_perceptron_tagger', quiet=True)
            nltk.download('punkt_tab', quiet=True)
        except:
            print("NLTK data already downloaded")
    
    def is_numbered_item(self, line):
        """Check if line is a numbered item"""
        return bool(re.match(r'^\d+[\./-]', line.strip()))
    
    def is_choice(self, line):
        """Check if line is a choice"""
        return bool(re.match(r'^[A-Ea-e][\s\(\./\-]', line.strip()))
    
    def has_multiple_choices(self, line):
        """Check if line contains multiple choices"""
        return bool(re.search(r'[A-E]\s*\([^)]+\)', line))
    
    def analyze_question(self, lines):
        main_question = None
        numbered_items = []
        choices = []
        
        # Get the question title (first line)
        if lines:
            main_question = lines[0].strip()
            
        # Process remaining lines
        for line in lines[1:]:
            line = line.strip()
            if not line or line.startswith("Réponse :"):
                continue
            
            # Check for numbered items
            if self.is_numbered_item(line):
                # Make sure we don't include the question title in numbered items
                if not line.startswith(main_question):
                    # Standardize numbered item format
                    line = re.sub(r'^\d+[\./-]\s*', lambda m: f"{m.group(0)[0]}- ", line)
                    numbered_items.append(line)
            # Check for multiple choices on one line
            elif self.has_multiple_choices(line):
                # Split multiple choices
                choices_matches = re.finditer(r'([A-E])\s*\(([^)]+)\)', line)
                for match in choices_matches:
                    letter, content = match.groups()
                    choices.append(f"{letter}. ({content})")
            # Check for single choice
            elif self.is_choice(line):
                # Keep the entire choice text
                if '.' not in line[:2]:  # If not already in correct format
                    line = re.sub(r'^([A-E])\s*[\(\./\-]\s*', r'\1. ', line)
                choices.append(line)
        
        return main_question, numbered_items, choices

def extract_qcm_questions(input_file):
    doc = Document(input_file)
    questions = []
    current_question = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        # Skip empty paragraphs
        if not text:
            continue
            
        # Check if this is the start of a new question
        if re.match(r'^(?:Q\d+|\d+\.)', text):
            if current_question:
                questions.append('\n'.join(current_question))
            current_question = [text]
        else:
            current_question.append(text)
    
    # Add the last question
    if current_question:
        questions.append('\n'.join(current_question))
    
    return questions

def find_highlighted_answers(paragraph):
    highlighted_answers = []
    for run in paragraph.runs:
        is_highlighted = (
            run.font.highlight_color or
            run.font.color.rgb or
            getattr(run.font, 'highlight', None)
        )
        
        if is_highlighted:
            text = run.text.strip()
            if text and text[0] in 'ABCDE':
                highlighted_answers.append(text[0])
    
    return ''.join(sorted(set(highlighted_answers)))

def process_qcm_document(input_file, output_file):
    doc = Document(input_file)
    output_doc = Document()
    
    # Set Calibri font, 11px and remove spacing
    style = output_doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = 139700
    style.paragraph_format.space_after = 0
    style.paragraph_format.space_before = 0
    
    current_question = []
    current_answers = []
    question_counter = 1
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        if not text:
            continue
            
        if re.match(r'^(?:Q\d+|\d+\.)', text):
            # Process previous question if exists
            if current_question and current_answers:  # Only process if we have answers
                write_question_block(output_doc, current_question, current_answers, question_counter)
                question_counter += 1
            
            current_question = [text]
            current_answers = []
        else:
            current_question.append(text)
            # Check for highlighted answers with our improved detection
            highlighted = find_highlighted_answers(paragraph)
            if highlighted:
                current_answers.extend(highlighted)
    
    # Process the last question
    if current_question and current_answers:  # Only process if we have answers
        write_question_block(output_doc, current_question, current_answers, question_counter)
    
    output_doc.save(output_file)
